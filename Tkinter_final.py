from tkinter import *
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


root = Tk()
# root.geometry('800x600')
root.title('Kitab Ghar Printing App')
Label(root, text = " Name ").grid(row = 0, sticky = W)
Label(root, text = " Phone Number ").grid(row = 1, sticky = W)
Label(root, text = " Address ").grid(row = 2, sticky = W)
Label(root, text = " Pin Code ").grid(row = 3, sticky = W)
Label(root, text = " City ").grid(row = 4, sticky = W)
Label(root, text = "State ").grid(row = 5, sticky = W)
Label(root, text = "Country ").grid(row = 6, sticky = W)

nm = Entry(root)
ph = Entry(root)
add = Entry(root)
pin= Entry(root)
city= Entry(root)
state = Entry(root)
country = Entry(root) 

nm.grid(row = 0, column = 1)
ph.grid(row = 1, column = 1)
add.grid(row = 2, column = 1)
pin.grid(row = 3, column = 1)
city.grid(row = 4, column = 1)
state.grid(row = 5, column = 1)
country.grid(row = 6, column = 1)


def getInput():

    a = nm.get()
    b = ph.get()
    c = add.get()
    d = pin.get()
    e = city.get()
    f = state.get()
    g = country.get()

    doc =  Document()
    doc.add_heading("Kitab Ghar ",0)
    thanks =doc.add_paragraph()

    thanks.add_run("thanks for ordering with kitab Ghar").bold = True
    doc_para = doc.add_paragraph()
    doc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc_para.add_run('TO\t\t').font.size = Pt(14)
    doc_para.add_run(a.rjust(24)).font.size = Pt(14)
    doc_para.add_run("\n").font.size = Pt(14)
    doc_para.add_run("Phone Number:").font.size = Pt(14)
    doc_para.add_run(b.rjust(15)).font.size = Pt(14)
    doc_para.add_run('\n').font.size = Pt(14)
    doc_para.add_run('address:\t\t').font.size = Pt(14)
    doc_para.add_run(c.rjust(16)).font.size = Pt(14)
    doc_para.add_run('\n').font.size = Pt(14)
    doc_para.add_run('pin:\t\t').font.size = Pt(14)
    doc_para.add_run(d.rjust(24)).font.size = Pt(14)
    doc_para.add_run('\n').font.size = Pt(14)
    doc_para.add_run("city:\t\t").font.size = Pt(14)
    doc_para.add_run(e.rjust(24)).font.size = Pt(14)
    doc_para.add_run('\n').font.size = Pt(14)
    doc_para.add_run("state:\t\t").font.size = Pt(14)
    doc_para.add_run(f.rjust(16)).font.size = Pt(14)
    doc_para.add_run("\nCountry:\t\t").font.size = Pt(14)
    doc_para.add_run(g.rjust(16)).font.size = Pt(14)

    kitab = doc.add_paragraph()
    kitab.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    kitab.add_run("From : name@ Online_Kitab_Ghar ")
    kitab.add_run("\n Address : VPO Singo(Bathinda)(151302) ")
    kitab.add_run(" \n Ph: 70092-61058 , 94656-87761")

    doc.save("ship.docx")
    
    os.startfile('ship.docx', 'print')

    
Button(root, text = "print",
           command = getInput).grid(row = 8, sticky = W)


mainloop()